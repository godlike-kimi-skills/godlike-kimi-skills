#!/usr/bin/env python3
"""
docx-skill - Word文档处理器

创建、编辑、格式化Word文档(.docx)，支持模板、表格、图片插入和文档合并

Usage:
    python main.py create --output document.docx --content "Hello World"
    python main.py create --output report.docx --input report.md
    python main.py template --template template.docx --output output.docx
    python main.py merge --input doc1.docx,doc2.docx --output merged.docx
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# 尝试导入 docx，如果不存在则给出友好提示
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxProcessor:
    """Word文档处理器"""
    
    def __init__(self):
        self.version = "1.0.0"
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )
    
    def create_document(
        self,
        output_path: str,
        content: Optional[str] = None,
        input_path: Optional[str] = None,
        styles: Optional[Dict] = None
    ) -> str:
        """
        创建新文档
        
        Args:
            output_path: 输出文件路径
            content: 文档内容（纯文本或JSON）
            input_path: 输入文件路径（Markdown、txt等）
            styles: 样式配置
            
        Returns:
            输出文件路径
        """
        doc = Document()
        
        # 设置默认样式
        self._apply_styles(doc, styles or {})
        
        if input_path:
            # 从文件读取内容
            self._add_content_from_file(doc, input_path)
        elif content:
            # 解析内容
            if content.strip().startswith('{'):
                # JSON格式
                try:
                    data = json.loads(content)
                    self._add_json_content(doc, data)
                except json.JSONDecodeError:
                    # 作为纯文本处理
                    self._add_text_content(doc, content)
            else:
                # 纯文本或Markdown
                self._add_text_content(doc, content)
        else:
            # 空文档
            doc.add_paragraph()
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def _add_content_from_file(self, doc: Document, input_path: str) -> None:
        """从文件添加内容"""
        path = Path(input_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        suffix = path.suffix.lower()
        
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if suffix == '.md':
            self._add_markdown_content(doc, content)
        elif suffix == '.json':
            data = json.loads(content)
            self._add_json_content(doc, data)
        else:
            # 纯文本
            self._add_text_content(doc, content)
    
    def _add_markdown_content(self, doc: Document, content: str) -> None:
        """添加 Markdown 内容"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # 标题
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
            # 列表
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', stripped):
                text = re.sub(r'^\d+\. ', '', stripped)
                doc.add_paragraph(text, style='List Number')
            # 表格
            elif stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
                # 解析表格
                table_lines = [stripped]
                i += 1  # 跳过分隔行
                while i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                    i += 1
                    table_lines.append(lines[i].strip())
                
                self._add_markdown_table(doc, table_lines)
            # 普通段落
            elif stripped:
                doc.add_paragraph(stripped)
            # 空行
            else:
                pass
            
            i += 1
    
    def _add_markdown_table(self, doc: Document, table_lines: List[str]) -> None:
        """添加 Markdown 表格"""
        if len(table_lines) < 2:
            return
        
        # 解析表头
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # 创建表格
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # 添加数据行
        for line in table_lines[2:]:  # 跳过分隔行
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            row_cells = table.add_row().cells
            for i, cell in enumerate(cells):
                if i < len(row_cells):
                    row_cells[i].text = cell
    
    def _add_json_content(self, doc: Document, data: Dict) -> None:
        """添加 JSON 格式的内容"""
        # 标题
        if 'title' in data:
            doc.add_heading(data['title'], level=1)
        
        # 段落
        if 'paragraphs' in data:
            for para in data['paragraphs']:
                if isinstance(para, str):
                    doc.add_paragraph(para)
                elif isinstance(para, dict):
                    text = para.get('text', '')
                    style = para.get('style', 'Normal')
                    p = doc.add_paragraph(text, style=style)
        
        # 章节
        if 'sections' in data:
            for section in data['sections']:
                if isinstance(section, dict):
                    heading = section.get('heading', '')
                    content = section.get('content', '')
                    
                    if heading:
                        doc.add_heading(heading, level=2)
                    if content:
                        doc.add_paragraph(content)
        
        # 表格
        if 'table' in data:
            table_data = data['table']
            self._add_table(doc, table_data)
        
        # 列表
        if 'list' in data:
            for item in data['list']:
                doc.add_paragraph(str(item), style='List Bullet')
    
    def _add_table(self, doc: Document, table_data: Dict) -> None:
        """添加表格"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers and not rows:
            return
        
        col_count = len(headers) if headers else len(rows[0]) if rows else 1
        table = doc.add_table(rows=1, cols=col_count)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header)
        
        # 数据行
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(cell)
    
    def _add_text_content(self, doc: Document, content: str) -> None:
        """添加纯文本内容"""
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    def _apply_styles(self, doc: Document, styles: Dict) -> None:
        """应用样式配置"""
        # 这里可以添加更多样式配置
        pass
    
    def edit_document(
        self,
        input_path: str,
        output_path: str,
        edits: Dict[str, Any]
    ) -> str:
        """
        编辑现有文档
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径
            edits: 编辑指令
            
        Returns:
            输出文件路径
        """
        doc = Document(input_path)
        
        # 文本替换
        if 'replace' in edits:
            replacements = edits['replace']
            for old_text, new_text in replacements.items():
                self._replace_text_in_doc(doc, old_text, new_text)
        
        # 添加内容
        if 'add' in edits:
            add_content = edits['add']
            if isinstance(add_content, list):
                for item in add_content:
                    if isinstance(item, dict):
                        if 'heading' in item:
                            doc.add_heading(item['heading'], level=item.get('level', 1))
                        if 'paragraph' in item:
                            doc.add_paragraph(item['paragraph'])
            elif isinstance(add_content, str):
                doc.add_paragraph(add_content)
        
        # 保存
        doc.save(output_path)
        return output_path
    
    def _replace_text_in_doc(self, doc: Document, old_text: str, new_text: str) -> None:
        """在文档中替换文本"""
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
    
    def merge_documents(self, input_paths: List[str], output_path: str) -> str:
        """
        合并多个文档
        
        Args:
            input_paths: 输入文件路径列表
            output_path: 输出文件路径
            
        Returns:
            输出文件路径
        """
        if not input_paths:
            raise ValueError("No input files provided")
        
        # 加载第一个文档
        merged_doc = Document(input_paths[0])
        
        # 添加后续文档
        for path in input_paths[1:]:
            doc = Document(path)
            
            # 添加分页符
            merged_doc.add_page_break()
            
            # 复制所有元素
            for element in doc.element.body:
                merged_doc.element.body.append(element)
        
        # 保存
        merged_doc.save(output_path)
        return output_path
    
    def use_template(
        self,
        template_path: str,
        output_path: str,
        variables: Optional[Dict] = None,
        data_file: Optional[str] = None
    ) -> List[str]:
        """
        使用模板生成文档
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径（或目录）
            variables: 变量字典
            data_file: 数据文件路径（批量生成）
            
        Returns:
            生成的文件路径列表
        """
        output_files = []
        
        if data_file and Path(data_file).exists():
            # 批量生成
            with open(data_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            # 创建输出目录
            output_dir = Path(output_path)
            if output_dir.suffix:  # 是文件路径
                output_dir = output_dir.parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # 简单的变量替换
                vars_dict = {'name': line, 'index': i + 1}
                if variables:
                    vars_dict.update(variables)
                
                # 生成文件名
                out_file = output_dir / f"document_{i+1:03d}.docx"
                
                # 使用模板生成
                self._process_template(template_path, str(out_file), vars_dict)
                output_files.append(str(out_file))
        
        else:
            # 单文件生成
            self._process_template(template_path, output_path, variables or {})
            output_files.append(output_path)
        
        return output_files
    
    def _process_template(
        self,
        template_path: str,
        output_path: str,
        variables: Dict
    ) -> None:
        """处理模板"""
        doc = Document(template_path)
        
        # 替换变量
        for paragraph in doc.paragraphs:
            for var_name, var_value in variables.items():
                placeholder = f"{{{{{var_name}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(var_value))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for var_name, var_value in variables.items():
                        placeholder = f"{{{{{var_name}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(var_value))
        
        doc.save(output_path)


def main():
    """主入口"""
    parser = argparse.ArgumentParser(
        description="Word文档处理器 - 创建、编辑、合并Word文档",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 创建文档
  python main.py create --output document.docx --content "Hello World"
  
  # 从Markdown创建
  python main.py create --output report.docx --input report.md
  
  # 使用模板
  python main.py template --template template.docx --output output.docx
  
  # 合并文档
  python main.py merge --input doc1.docx,doc2.docx --output merged.docx
        """
    )
    
    subparsers = parser.add_subparsers(dest="action", help="可用操作")
    
    # create 命令
    create_parser = subparsers.add_parser("create", help="创建新文档")
    create_parser.add_argument("--output", "-o", required=True, help="输出文件路径")
    create_parser.add_argument("--input", "-i", help="输入文件路径")
    create_parser.add_argument("--content", "-c", help="文档内容")
    create_parser.add_argument("--styles", "-s", help="样式配置（JSON）")
    
    # edit 命令
    edit_parser = subparsers.add_parser("edit", help="编辑文档")
    edit_parser.add_argument("--input", "-i", required=True, help="输入文件路径")
    edit_parser.add_argument("--output", "-o", required=True, help="输出文件路径")
    edit_parser.add_argument("--edits", "-e", required=True, help="编辑指令（JSON）")
    
    # merge 命令
    merge_parser = subparsers.add_parser("merge", help="合并文档")
    merge_parser.add_argument("--input", "-i", required=True, help="输入文件路径（逗号分隔）")
    merge_parser.add_argument("--output", "-o", required=True, help="输出文件路径")
    
    # template 命令
    template_parser = subparsers.add_parser("template", help="使用模板")
    template_parser.add_argument("--template", "-t", required=True, help="模板文件路径")
    template_parser.add_argument("--output", "-o", required=True, help="输出文件路径")
    template_parser.add_argument("--variables", "-v", help="变量（JSON）")
    template_parser.add_argument("--data-file", "-d", help="数据文件（批量生成）")
    
    parser.add_argument("--version", "-V", action="version", version="%(prog)s 1.0.0")
    
    args = parser.parse_args()
    
    if not args.action:
        parser.print_help()
        return 1
    
    # 检查依赖
    if not DOCX_AVAILABLE:
        print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
        return 1
    
    try:
        processor = DocxProcessor()
        
        if args.action == "create":
            result = processor.create_document(
                output_path=args.output,
                content=args.content,
                input_path=args.input,
                styles=json.loads(args.styles) if args.styles else None
            )
            print(f"✅ Document created: {result}")
        
        elif args.action == "edit":
            edits = json.loads(args.edits)
            result = processor.edit_document(
                input_path=args.input,
                output_path=args.output,
                edits=edits
            )
            print(f"✅ Document edited: {result}")
        
        elif args.action == "merge":
            input_paths = [p.strip() for p in args.input.split(",")]
            result = processor.merge_documents(input_paths, args.output)
            print(f"✅ Documents merged: {result}")
        
        elif args.action == "template":
            variables = json.loads(args.variables) if args.variables else None
            results = processor.use_template(
                template_path=args.template,
                output_path=args.output,
                variables=variables,
                data_file=args.data_file
            )
            print(f"✅ Template processed: {len(results)} file(s) generated")
            for f in results:
                print(f"   - {f}")
        
        return 0
        
    except Exception as e:
        print(f"❌ Error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
